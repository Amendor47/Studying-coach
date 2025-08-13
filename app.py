from flask import Flask, jsonify, request, render_template, Response, send_file
import io
import csv
import os
import sys
import time
import json
import logging
from datetime import datetime

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('studying_coach.log') if os.path.exists('.') else logging.NullHandler()
    ]
)
logger = logging.getLogger(__name__)

from services.analyzer import analyze_offline
from services.heuristics import ai_needed, readability, density
from services.store import load_db, save_db
from typing import Dict
from services.validate import validate_items, seed_seen_hashes
from services.ai import analyze_text
from services.rag import get_context
from services.planner import generate_plan
from services.scheduler import (
    due_cards,
    update_srs,
    build_exercises_from_cards,
    save_progress,
)
from services.parsers import extract_text
from services.webfetch import web_context_from_query
from pathlib import Path
from services.teacher import LocalTeacher
from services.config import load_settings
from services.llm_adapter import LLMClient
from services.performance_cache import performance_cache, cached
from services.local_llm import get_local_llm
from services.ai_pipeline import get_ai_pipeline

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

BASE_DIR = getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))
app = Flask(
    __name__,
    static_folder=os.path.join(BASE_DIR, "static"),
    template_folder=os.path.join(BASE_DIR, "templates"),
)

# Configure CORS for development
from flask_cors import CORS
cors_origins = os.getenv("CORS_ORIGINS", "").split(",") if os.getenv("CORS_ORIGINS") else []
if cors_origins:
    CORS(app, origins=cors_origins)
else:
    # Enable permissive CORS for development
    CORS(app, resources={r"/*": {"origins": "*"}})

teacher = LocalTeacher()


# ---- configuration & API key management ----
ENV_FILE = Path(".env")


@app.route("/api/config", methods=["GET", "POST"])
def api_config():
    """Expose presence of OPENAI_API_KEY and allow setting it."""
    if request.method == "GET":
        # Cache configuration check
        cache_key = "api_config"
        cached_result = performance_cache.get("config", cache_key)
        if cached_result:
            return jsonify(cached_result)
        
        result = {"has_key": bool(os.getenv("OPENAI_API_KEY"))}
        performance_cache.set("config", cache_key, result, ttl_seconds=60)
        return jsonify(result)

    data = request.get_json(force=True)
    key = data.get("key", "").strip()
    if not key:
        return jsonify({"saved": False}), 400
    os.environ["OPENAI_API_KEY"] = key
    ENV_FILE.write_text(f"OPENAI_API_KEY={key}\n", encoding="utf-8")
    
    # Invalidate config cache
    performance_cache.invalidate("config", "api_config")
    
    return jsonify({"saved": True})


@app.route("/api/health/llm")
@cached("health", ttl_seconds=300)  # Cache for 5 minutes
def health_llm():
    """Check availability of the configured LLM provider."""
    settings = load_settings()
    client = LLMClient.from_settings(settings)
    ok, msg = client.healthcheck()
    return jsonify({"provider": settings.provider, "model": settings.model, "ok": ok, "msg": msg})


@app.route("/api/upload", methods=["POST"])
def upload_file():
    """Handle document upload and full analysis pipeline."""
    uploaded = request.files.get("file")
    if not uploaded:
        return jsonify({"error": "no file"}), 400
    use_ai = request.form.get("use_ai", "false").lower() == "true"
    use_advanced_analysis = request.form.get("use_advanced", "false").lower() == "true"
    minutes = int(request.form.get("session_minutes", 0))
    from datetime import date
    from services.tempfiles import safe_save_upload, safe_unlink

    suffix = Path(uploaded.filename).suffix or ".bin"
    tmp_path = safe_save_upload(uploaded, suffix)
    
    try:
        # Use advanced document analysis if requested
        if use_advanced_analysis:
            from services.advanced_document_analysis import advanced_document_analyzer
            
            analysis_result = advanced_document_analyzer.analyze_document(
                tmp_path.as_posix(), uploaded.filename
            )
            
            # Convert analyzed content back to text for existing pipeline
            segments = analysis_result.get("segments", [])
            text = "\n\n".join(seg["content"] for seg in segments if seg["content"])
            
            # Store advanced analysis results for later use (converted to serializable format)
            advanced_metadata = {}
            if analysis_result:
                doc_meta = analysis_result.get("metadata")
                if doc_meta:
                    # Convert DocumentMetadata to dict
                    advanced_metadata = {
                        "document_metadata": {
                            "filename": doc_meta.filename,
                            "file_type": doc_meta.file_type,
                            "document_type": doc_meta.document_type.value if hasattr(doc_meta.document_type, 'value') else str(doc_meta.document_type),
                            "language": doc_meta.language,
                            "page_count": doc_meta.page_count,
                            "word_count": doc_meta.word_count,
                            "has_formulas": doc_meta.has_formulas,
                            "has_images": doc_meta.has_images,
                            "has_tables": doc_meta.has_tables,
                            "complexity_score": doc_meta.complexity_score,
                            "reading_level": doc_meta.reading_level
                        },
                        "formulas": [
                            {
                                "original_text": f.original_text,
                                "latex_representation": f.latex_representation,
                                "formula_type": f.formula_type,
                                "variables": f.variables,
                                "constants": f.constants,
                                "operations": f.operations,
                                "complexity_level": f.complexity_level,
                                "subject_area": f.subject_area
                            }
                            for f in analysis_result.get("formulas", [])
                        ],
                        "classified_content": analysis_result.get("classified_content", {}),
                        "learning_objectives": analysis_result.get("learning_objectives", []),
                        "concept_map": analysis_result.get("concept_map", {}),
                        "educational_insights": analysis_result.get("educational_insights", {})
                    }
        else:
            # Fallback to original text extraction
            text = extract_text(tmp_path.as_posix(), uploaded.filename)
            advanced_metadata = {}
            
    finally:
        safe_unlink(tmp_path)

    db = load_db()
    seed_items = [{"kind": "card", "payload": c} for c in db.get("cards", [])]
    seed_items += [{"kind": "exercise", "payload": e} for e in db.get("exercises", [])]
    seed_seen_hashes(seed_items)

    offline_drafts = analyze_offline(text)
    drafts = offline_drafts
    if use_ai and os.getenv("OPENAI_API_KEY"):
        context = get_context(text)
        combined = "\n".join(context) + "\n\n" + text if context else text
        ai_drafts = analyze_text(combined, "upload")
        drafts = validate_items(offline_drafts + ai_drafts)
    else:
        drafts = validate_items(offline_drafts)

    for d in drafts:
        d.setdefault("status", "new")
        
        # Enrich with advanced analysis metadata if available
        if advanced_metadata:
            d.setdefault("advanced_metadata", {})
            d["advanced_metadata"].update(advanced_metadata)

    db["drafts"].extend(drafts)
    for d in drafts:
        payload = d.get("payload", {})
        if d.get("kind") == "card":
            payload.setdefault("id", d.get("id"))
            payload.setdefault(
                "srs",
                {"EF": 2.5, "interval": 1, "reps": 0, "due": date.today().isoformat()},
            )
            db.setdefault("cards", []).append(payload)
        elif d.get("kind") == "exercise":
            db.setdefault("exercises", []).append(payload)
        elif d.get("kind") == "course":
            db.setdefault("courses", []).append(payload)
    save_db(db)

    plan = generate_plan(db.get("drafts", []))
    themes = list({d.get("payload", {}).get("theme", "Général") for d in drafts})
    due = due_cards(db)
    
    response_data = {
        "saved": len(drafts),
        "plan": plan,
        "themes": themes,
        "due": due,
        "minutes": minutes,
    }
    
    # Include advanced analysis results if available
    if advanced_metadata:
        response_data["advanced_analysis"] = {
            "document_type": advanced_metadata.get("document_metadata", {}).get("document_type"),
            "complexity_score": advanced_metadata.get("document_metadata", {}).get("complexity_score"),
            "reading_level": advanced_metadata.get("document_metadata", {}).get("reading_level"),
            "formula_count": len(advanced_metadata.get("formulas", [])),
            "learning_objectives": advanced_metadata.get("learning_objectives", []),
            "estimated_study_time": advanced_metadata.get("educational_insights", {}).get("estimated_study_time", 0),
            "recommendations": advanced_metadata.get("educational_insights", {}).get("learning_recommendations", [])
        }
    
    return jsonify(response_data)


@app.route("/api/offline/analyze", methods=["POST"])
def offline_analyze():
    data = request.get_json(force=True)
    text = data.get("text", "")
    
    # Create cache key from text hash
    import hashlib
    text_hash = hashlib.md5(text.encode()).hexdigest()
    cache_key = f"offline_analysis:{text_hash}"
    
    # Try cache first
    cached_result = performance_cache.get("analysis", cache_key)
    if cached_result:
        # Record cache hit pattern
        performance_cache.record_pattern("default", "offline_analyze", [cache_key])
        return jsonify(cached_result)
    
    # Perform analysis
    drafts_raw = analyze_offline(text)
    drafts = validate_items(drafts_raw)
    need_ai = ai_needed(text, drafts)
    meta = {"readability": readability(text), "density": density(text, drafts)}
    
    result = {"drafts": drafts, "need_ai": need_ai, "meta": meta}
    
    # Cache result
    performance_cache.set("analysis", cache_key, result, ttl_seconds=1800)  # 30 minutes
    
    return jsonify(result)


@app.route("/api/ai/analyze", methods=["POST"])
def ai_analyze():
    """Legacy AI analyze - enhanced with new pipeline capabilities"""
    data = request.get_json(force=True)
    text = data.get("text", "")
    force = data.get("force", False)
    reason = data.get("reason", "")
    goals = data.get("goals", [])
    
    # If enhanced analysis is requested, use new pipeline
    if goals or data.get("enhanced", False):
        try:
            pipeline = get_ai_pipeline()
            analysis = pipeline.analyze_text(text, goals)
            return jsonify({
                "drafts": [{"title": kp, "content": kp} for kp in analysis.key_points],
                "reason": "enhanced_pipeline",
                "enhanced": analysis.to_dict()
            })
        except Exception as e:
            # Fall back to legacy analysis
            pass
    
    # Legacy analysis path
    offline_raw = analyze_offline(text)
    offline_drafts = validate_items(offline_raw)
    need_ai = ai_needed(text, offline_drafts)
    if not (need_ai or force):
        return jsonify({"drafts": offline_drafts, "reason": "offline"})
    reason = reason or ("heuristic" if need_ai else "forced")
    context = get_context(text)
    combined = "\n".join(context) + "\n\n" + text if context else text
    ai_drafts = analyze_text(combined, reason)
    drafts = validate_items(offline_drafts + ai_drafts)
    return jsonify({"drafts": drafts, "reason": reason})


@app.route("/api/improve", methods=["POST"])
def improve_text():
    """Simple text improvement endpoint using LLM"""
    try:
        data = request.get_json()
        logger.info(f"Improve endpoint called with data: {data}")
        
        if not data or 'text' not in data:
            logger.warning("Missing 'text' parameter in improve request")
            return jsonify({"error": "Missing 'text' parameter"}), 400
        
        text = data.get('text', '')
        model = data.get('model', 'llama3:8b')
        
        if not text.strip():
            logger.warning("Empty text provided to improve endpoint")
            return jsonify({"error": "Empty text provided"}), 400
        
        # Try to use local LLM first
        try:
            llm = get_local_llm()
            result = llm.improve_text(text, model=model)
            if result:
                logger.info(f"Successfully improved text using local LLM")
                return jsonify({
                    "ok": True,
                    "result": result,
                    "provider": "local",
                    "model": model
                })
        except Exception as e:
            logger.error(f"Local LLM failed in improve endpoint: {e}")
        
        # Fallback to basic improvement
        improved_text = f"Improved version of: {text}"
        logger.info(f"Using fallback improvement for text")
        return jsonify({
            "ok": True,
            "result": improved_text,
            "provider": "fallback",
            "model": "none"
        })
        
    except Exception as e:
        logger.error(f"Error in improve endpoint: {e}", exc_info=True)
        return jsonify({
            "ok": False,
            "error": str(e)
        }), 500


@app.route("/api/save", methods=["POST"])
def save_items():
    data = request.get_json(force=True)
    items = data.get("items", [])
    db = load_db()
    seed_items = [{"kind": "card", "payload": c} for c in db.get("cards", [])]
    seed_items += [{"kind": "exercise", "payload": e} for e in db.get("exercises", [])]
    seed_seen_hashes(seed_items)

    validated = validate_items(items)
    for d in validated:
        d.setdefault("status", "new")
    db["drafts"].extend(validated)
    # also populate dedicated card/exercise stores with SRS defaults
    from datetime import date

    for d in validated:
        payload = d.get("payload", {})
        if d.get("kind") == "card":
            payload.setdefault("id", d.get("id"))
            payload.setdefault(
                "srs",
                {"EF": 2.5, "interval": 1, "reps": 0, "due": date.today().isoformat()},
            )
            db.setdefault("cards", []).append(payload)
        elif d.get("kind") == "exercise":
            db.setdefault("exercises", []).append(payload)
    save_db(db)
    return jsonify({"saved": len(validated)})


@app.route("/api/plan", methods=["GET"])
def plan_route():
    db = load_db()
    plan = generate_plan(db.get("drafts", []))
    return jsonify({"plan": plan})


@app.route("/api/due", methods=["GET"])
def due_route():
    db = load_db()
    cards = due_cards(db)
    exercises = db.get("exercises", []) + build_exercises_from_cards(cards)
    return jsonify({"cards": cards, "exercises": exercises})


@app.route("/api/review/<cid>", methods=["POST"])
def review_route(cid: str):
    data = request.get_json(force=True)
    quality = int(data.get("quality", 0))
    db = load_db()
    for c in db.get("cards", []):
        if c.get("id") == cid:
            update_srs(c, quality)
            break
    save_progress(db)
    return jsonify({"ok": True})


@app.route("/api/themes", methods=["GET"])
def themes_route():
    """Return available themes and counts."""
    db = load_db()
    counts: Dict[str, Dict[str, int]] = {}
    for d in db.get("drafts", []):
        theme = d.get("payload", {}).get("theme", "Général")
        entry = counts.setdefault(theme, {"cards": 0, "exercises": 0})
        if d.get("kind") == "card":
            entry["cards"] += 1
        elif d.get("kind") == "exercise":
            entry["exercises"] += 1
    themes = [
        {"name": t, "cards": v["cards"], "exercises": v["exercises"]}
        for t, v in sorted(counts.items())
    ]
    return jsonify({"themes": themes})


@app.route("/api/fiches/<theme>", methods=["GET"])
def fiches_by_theme(theme: str):
    db = load_db()
    items = [
        d
        for d in db.get("drafts", [])
        if d.get("payload", {}).get("theme", "Général") == theme
    ]
    cards = [d for d in items if d.get("kind") == "card"]
    exercises = [d for d in items if d.get("kind") == "exercise"]
    courses = [d for d in items if d.get("kind") == "course"]
    return jsonify({"cards": cards, "exercises": exercises, "courses": courses})


@app.route("/api/card/<cid>/status", methods=["POST"])
def update_card_status(cid: str):
    data = request.get_json(force=True)
    status = data.get("status", "")
    db = load_db()
    for d in db.get("drafts", []):
        if d.get("id") == cid:
            d["status"] = status
            break
    save_db(db)
    return jsonify({"ok": True})


@app.route("/api/export/<fmt>", methods=["GET"])
def export_route(fmt: str):
    db = load_db()
    cards = db.get("cards", [])
    if fmt == "csv":
        buf = io.StringIO()
        writer = csv.writer(buf, delimiter=";")
        writer.writerow(["front", "back", "tags"])
        for c in cards:
            writer.writerow([c.get("front"), c.get("back"), c.get("theme", "")])
        buf.seek(0)
        return Response(
            buf.getvalue(),
            mimetype="text/csv",
            headers={"Content-Disposition": "attachment; filename=fiches.csv"},
        )
    if fmt == "pdf":
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            buf = io.BytesIO()
            c = canvas.Canvas(buf, pagesize=letter)
            y = 750
            for card in cards:
                c.drawString(72, y, f"{card.get('front')}: {card.get('back')}")
                y -= 20
                if y < 72:
                    c.showPage()
                    y = 750
            c.save()
            buf.seek(0)
            return send_file(buf, as_attachment=True, download_name="fiches.pdf", mimetype="application/pdf")
        except Exception:
            return jsonify({"error": "pdf export unavailable"}), 500
    if fmt == "docx":
        try:
            from docx import Document
            buf = io.BytesIO()
            doc = Document()
            for card in cards:
                doc.add_heading(card.get("front"), level=2)
                doc.add_paragraph(card.get("back"))
            doc.save(buf)
            buf.seek(0)
            return send_file(buf, as_attachment=True, download_name="fiches.docx")
        except Exception:
            return jsonify({"error": "docx export unavailable"}), 500
    return jsonify({"error": "format inconnu"}), 400


@app.route("/api/web/search", methods=["GET"])
def web_search():
    q = request.args.get("q", "").strip()
    if not q:
        return jsonify({"results": []})
    try:
        results = web_context_from_query(q, k_pages=5)
    except Exception:
        results = []
    return jsonify({
        "results": [
            {"title": r["title"], "url": r["url"], "excerpt": r["excerpt"][:300]}
            for r in results
        ]
    })


@app.route("/api/web/enrich", methods=["POST"])
def web_enrich():
    data = request.get_json(force=True)
    query = data.get("query", "").strip()
    use_ai = bool(data.get("use_ai", False))
    if not query:
        return jsonify({"added": 0, "drafts": [], "citations": []})

    try:
        pages = web_context_from_query(query, k_pages=3, max_chars=5000)
    except Exception:
        pages = []
    combined = "\n\n".join([p["excerpt"] for p in pages])

    db = load_db()
    seed_items = [{"kind": "card", "payload": c} for c in db.get("cards", [])]
    seed_items += [{"kind": "exercise", "payload": e} for e in db.get("exercises", [])]
    seed_seen_hashes(seed_items)

    offline_drafts = analyze_offline(combined)

    drafts = offline_drafts
    if use_ai and os.getenv("OPENAI_API_KEY"):
        context_text = "".join(f"\n# {p['title']}\n{p['excerpt']}\n" for p in pages)
        ai_drafts = analyze_text(context_text, reason="web+rag")
        drafts = validate_items(offline_drafts + ai_drafts)
    else:
        drafts = validate_items(offline_drafts)
    from datetime import date
    for d in drafts:
        d.setdefault("status", "new")
        db["drafts"].append(d)
        payload = d.get("payload", {})
        if d.get("kind") == "card":
            payload.setdefault("id", d.get("id"))
            payload.setdefault(
                "srs",
                {"EF": 2.5, "interval": 1, "reps": 0, "due": date.today().isoformat()},
            )
            db.setdefault("cards", []).append(payload)
        elif d.get("kind") == "exercise":
            db.setdefault("exercises", []).append(payload)
        elif d.get("kind") == "course":
            db.setdefault("courses", []).append(payload)
    save_db(db)

    citations = [{"title": p["title"], "url": p["url"]} for p in pages]
    return jsonify({"added": len(drafts), "drafts": drafts, "citations": citations})


@app.route("/api/chat", methods=["POST"])
def chat_route():
    data = request.get_json(force=True)
    msg = data.get("message", "")
    answer = teacher.chat(msg)
    return jsonify({"answer": answer})


@app.route("/api/advanced/analyze_document", methods=["POST"])
def advanced_document_analysis():
    """Advanced document analysis with educational intelligence"""
    uploaded = request.files.get("file")
    if not uploaded:
        return jsonify({"error": "no file"}), 400
        
    from services.tempfiles import safe_save_upload, safe_unlink
    from services.advanced_document_analysis import advanced_document_analyzer
    
    suffix = Path(uploaded.filename).suffix or ".bin"
    tmp_path = safe_save_upload(uploaded, suffix)
    
    try:
        analysis_result = advanced_document_analyzer.analyze_document(
            tmp_path.as_posix(), uploaded.filename
        )
        return jsonify(analysis_result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        safe_unlink(tmp_path)


@app.route("/api/advanced/educational_content", methods=["POST"])
def generate_educational_content():
    """Generate adaptive educational content using advanced AI"""
    data = request.get_json(force=True)
    text = data.get("text", "")
    user_id = data.get("user_id", "default")
    learning_objective = data.get("learning_objective", {})
    
    if not text:
        return jsonify({"error": "No text provided"}), 400
        
    try:
        from services.educational_ai import educational_ai
        from services.contextual_memory import contextual_memory
        
        # Get user profile for personalization
        user_profile = contextual_memory.get_or_create_profile(user_id)
        
        # Create learning objective object
        from services.educational_ai import LearningObjective, DifficultyLevel
        objective = LearningObjective(
            id=learning_objective.get("id", "default"),
            title=learning_objective.get("title", "General Learning"),
            description=learning_objective.get("description", "Learn the provided content"),
            level=DifficultyLevel(learning_objective.get("level", 2)),  # Default to intermediate
            prerequisites=learning_objective.get("prerequisites", []),
            mastery_criteria=learning_objective.get("mastery_criteria", [])
        )
        
        # Generate adaptive content
        content = educational_ai.generate_adaptive_content(text, user_profile, objective)
        
        return jsonify(content)
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/advanced/learning_analytics", methods=["GET"])
def learning_analytics():
    """Get comprehensive learning analytics for a user"""
    user_id = request.args.get("user_id", "default")
    
    try:
        from services.contextual_memory import contextual_memory
        analytics = contextual_memory.generate_learning_analytics(user_id)
        return jsonify(analytics)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/advanced/learning_interaction", methods=["POST"])
def record_learning_interaction():
    """Record a learning interaction for memory tracking"""
    data = request.get_json(force=True)
    
    user_id = data.get("user_id", "default")
    concept_id = data.get("concept_id", "")
    concept_name = data.get("concept_name", "")
    is_correct = data.get("is_correct", False)
    response_time = data.get("response_time", 0.0)
    confidence = data.get("confidence", 0.5)
    context = data.get("context", "")
    
    if not concept_id or not concept_name:
        return jsonify({"error": "concept_id and concept_name required"}), 400
    
    try:
        from services.contextual_memory import contextual_memory
        contextual_memory.record_learning_interaction(
            user_id, concept_id, concept_name, is_correct, 
            response_time, confidence, context
        )
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/advanced/due_concepts", methods=["GET"])
def get_due_concepts():
    """Get concepts due for review with intelligent prioritization"""
    user_id = request.args.get("user_id", "default")
    max_count = int(request.args.get("max_count", 10))
    
    try:
        from services.contextual_memory import contextual_memory
        due_concepts = contextual_memory.get_due_concepts(user_id, max_count)
        
        # Convert to JSON-serializable format
        result = []
        for concept in due_concepts:
            result.append({
                "concept_id": concept.concept_id,
                "concept_name": concept.concept_name,
                "knowledge_state": concept.knowledge_state.value,
                "memory_strength": concept.memory_strength,
                "confidence_level": concept.confidence_level,
                "accuracy_rate": concept.accuracy_rate(),
                "days_overdue": (datetime.now() - concept.next_review).days,
                "review_count": concept.review_count
            })
        
        return jsonify({"due_concepts": result})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/advanced/learning_path", methods=["POST"])
def optimize_learning_path():
    """Generate optimized learning path"""
    data = request.get_json(force=True)
    user_id = data.get("user_id", "default")
    target_concepts = data.get("target_concepts", [])
    
    if not target_concepts:
        return jsonify({"error": "target_concepts required"}), 400
    
    try:
        from services.contextual_memory import contextual_memory
        learning_path = contextual_memory.optimize_learning_path(user_id, target_concepts)
        return jsonify({"learning_path": learning_path})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/advanced/rag_search", methods=["POST"])
def advanced_rag_search():
    """Perform advanced RAG search with educational context"""
    data = request.get_json(force=True)
    query = data.get("query", "")
    include_levels = data.get("include_levels", ["paragraph", "section"])
    top_k = int(data.get("top_k", 5))
    
    if not query:
        return jsonify({"error": "query required"}), 400
    
    try:
        from services.advanced_rag import advanced_rag, ChunkType
        
        # Convert string levels to enum
        level_map = {
            "document": ChunkType.DOCUMENT,
            "chapter": ChunkType.CHAPTER,
            "section": ChunkType.SECTION,
            "paragraph": ChunkType.PARAGRAPH,
            "sentence": ChunkType.SENTENCE
        }
        
        chunk_types = [level_map[level] for level in include_levels if level in level_map]
        
        results = advanced_rag.multi_level_retrieval(query, top_k, chunk_types)
        
        # Convert results to JSON-serializable format
        serialized_results = []
        for result in results:
            serialized_results.append({
                "chunk_id": result.chunk.id,
                "content": result.chunk.content,
                "chunk_type": result.chunk.chunk_type.value,
                "content_type": result.chunk.content_type.value,
                "relevance_score": result.relevance_score,
                "educational_value": result.educational_value,
                "concept_coverage": result.concept_coverage,
                "citations": result.citations,
                "confidence": result.confidence
            })
        
        return jsonify({"results": serialized_results})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/advanced/concept_network", methods=["GET"])
def get_concept_network():
    """Get concept network for knowledge graph visualization"""
    concept = request.args.get("concept", "")
    depth = int(request.args.get("depth", 2))
    
    if not concept:
        return jsonify({"error": "concept parameter required"}), 400
    
    # Cache concept networks as they're expensive to compute
    cache_key = f"concept_network:{concept}:{depth}"
    cached_result = performance_cache.get("concept_network", cache_key)
    if cached_result:
        return jsonify(cached_result)
    
    try:
        from services.advanced_rag import advanced_rag
        network = advanced_rag.get_concept_network(concept, depth)
        
        # Cache for 1 hour
        performance_cache.set("concept_network", cache_key, network, ttl_seconds=3600)
        
        return jsonify(network)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/performance/stats", methods=["GET"])
def performance_stats():
    """Get comprehensive performance statistics"""
    stats = performance_cache.get_stats()
    
    # Add application-specific metrics
    app_stats = {
        "flask": {
            "request_count": getattr(app, '_request_count', 0),
            "error_count": getattr(app, '_error_count', 0),
        },
        "database": {
            "connection_pool_size": 10,  # Would be actual pool size
            "active_connections": 2,     # Would be actual active connections
        }
    }
    
    return jsonify({
        "cache": stats,
        "application": app_stats,
        "timestamp": datetime.now().isoformat()
    })


@app.route("/api/performance/clear_cache", methods=["POST"])
def clear_cache():
    """Clear performance caches"""
    success = performance_cache.clear_all()
    return jsonify({
        "success": success,
        "message": "All caches cleared" if success else "Failed to clear some caches"
    })


@app.route("/api/performance/preload", methods=["POST"])
def preload_cache():
    """Preload cache based on usage patterns"""
    data = request.get_json(force=True)
    user_id = data.get("user_id", "default")
    current_action = data.get("current_action", "")
    
    # Predict next resources
    predictions = performance_cache.predict_next_resources(user_id, current_action)
    
    # Preload predicted resources
    preload_patterns = []
    for resource in predictions[:5]:  # Top 5 predictions
        preload_patterns.append(("analysis", resource, {}))
    
    performance_cache.preload(preload_patterns)
    
    return jsonify({
        "predictions": predictions,
        "preloaded": len(preload_patterns)
    })


# Performance monitoring middleware
@app.before_request
def before_request():
    request._start_time = time.perf_counter()
    # Track request count
    app._request_count = getattr(app, '_request_count', 0) + 1


@app.after_request
def after_request(response):
    """Safe after_request handler that prevents RuntimeError on static files."""
    # Calculate response time
    if hasattr(request, '_start_time'):
        response_time = (time.perf_counter() - request._start_time) * 1000
        response.headers['X-Response-Time'] = f"{response_time:.2f}ms"
        
        # Log slow requests
        if response_time > 1000:  # > 1 second
            print(f"Slow request: {request.endpoint} took {response_time:.2f}ms")
    
    # Skip ETag generation for static files
    if request.endpoint == 'static':
        # Add cache headers for static content without accessing response.data
        response.headers['Cache-Control'] = 'public, max-age=31536000'  # 1 year
    else:
        # Skip for streaming responses
        if response.is_streamed:
            pass
        # Skip for file downloads and binary responses
        elif response.mimetype and (
            response.mimetype.startswith('image/') or 
            response.mimetype.startswith('video/') or
            response.mimetype.startswith('audio/') or
            response.mimetype == 'application/octet-stream'
        ):
            pass
        else:
            # Only add ETag for regular responses with accessible data
            try:
                if hasattr(response, 'get_data'):
                    data = response.get_data(as_text=False)
                    if data:
                        response.headers['ETag'] = f'"{hash(data)}"'
            except (RuntimeError, AttributeError):
                # Silently skip ETag generation if data is not accessible
                pass
    
    # Add security headers
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY' 
    response.headers['X-XSS-Protection'] = '1; mode=block'
    
    return response


@app.errorhandler(500)
def internal_error(error):
    app._error_count = getattr(app, '_error_count', 0) + 1
    return jsonify({"error": "Internal server error"}), 500


# ===== NEW API ENDPOINTS FOR DRAG & DROP AND LLM INTEGRATION =====

@app.route("/api/flashcards/reorder", methods=["POST"])
def reorder_flashcards():
    """Persist flashcard reorder from drag & drop"""
    try:
        data = request.get_json()
        deck_id = data.get("deck_id")
        order = data.get("order", [])
        
        if not deck_id or not order:
            return jsonify({"error": "Missing deck_id or order"}), 400
        
        # Load current database
        db = load_db()
        
        # Find cards for this deck and reorder them
        cards = db.get("cards", [])
        deck_cards = [c for c in cards if c.get("theme") == deck_id]
        other_cards = [c for c in cards if c.get("theme") != deck_id]
        
        # Create a mapping of card IDs to cards
        card_map = {c.get("id", str(i)): c for i, c in enumerate(deck_cards)}
        
        # Reorder according to the new order
        reordered_cards = []
        for card_id in order:
            if card_id in card_map:
                reordered_cards.append(card_map[card_id])
        
        # Add any missing cards (in case of sync issues)
        included_ids = set(order)
        for card in deck_cards:
            card_id = card.get("id", str(len(reordered_cards)))
            if card_id not in included_ids:
                reordered_cards.append(card)
        
        # Update database with new order
        db["cards"] = other_cards + reordered_cards
        save_db(db)
        
        return jsonify({
            "success": True,
            "deck_id": deck_id,
            "reordered_count": len(reordered_cards)
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/generate", methods=["POST"])
def ai_generate():
    """Generate text using local LLM with optional streaming"""
    try:
        data = request.get_json()
        prompt = data.get("prompt", "")
        system = data.get("system")
        temperature = data.get("temperature", 0.2)
        stream = data.get("stream", False)
        
        if not prompt:
            return jsonify({"error": "Missing prompt"}), 400
        
        llm = get_local_llm()
        
        if stream:
            # Return Server-Sent Events for streaming
            def generate_stream():
                try:
                    response_stream = llm.generate(
                        prompt=prompt,
                        system=system,
                        temperature=temperature,
                        stream=True
                    )
                    
                    for chunk in response_stream:
                        if hasattr(chunk, "get"):
                            content = chunk.get("message", {}).get("content", "")
                        else:
                            content = str(chunk)
                        
                        if content:
                            yield f"data: {json.dumps({'content': content})}\n\n"
                    
                    yield f"data: {json.dumps({'done': True})}\n\n"
                    
                except Exception as e:
                    yield f"data: {json.dumps({'error': str(e)})}\n\n"
            
            return Response(
                generate_stream(),
                mimetype='text/event-stream',
                headers={
                    'Cache-Control': 'no-cache',
                    'Connection': 'keep-alive',
                    'Access-Control-Allow-Origin': '*'
                }
            )
        else:
            # Regular JSON response
            response = llm.generate(
                prompt=prompt,
                system=system,
                temperature=temperature,
                stream=False
            )
            
            return jsonify({
                "text": response.get("text", ""),
                "model": llm.config.model,
                "provider": llm.config.provider
            })
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai/exercises", methods=["POST"])
def ai_exercises():
    """Generate educational exercises from text"""
    try:
        data = request.get_json()
        text = data.get("text", "")
        
        if not text:
            return jsonify({"error": "Missing text"}), 400
        
        # Use AI pipeline to generate study materials
        pipeline = get_ai_pipeline()
        materials = pipeline.generate_study_materials(text)
        
        # Extract exercises from the materials
        exercises = []
        
        # Add flashcards as recall exercises
        for card in materials.flashcards:
            exercises.append({
                "type": "flashcard",
                "question": card.get("front", ""),
                "answer": card.get("back", ""),
                "difficulty": "medium"
            })
        
        # Add quizzes
        exercises.extend(materials.quizzes)
        
        # Add mnemonics as memory exercises
        for mnemonic in materials.mnemonics:
            exercises.append({
                "type": "mnemonic",
                "concept": mnemonic.get("concept", ""),
                "device": mnemonic.get("mnemonic", ""),
                "difficulty": "easy"
            })
        
        return jsonify({
            "exercises": exercises,
            "total_count": len(exercises),
            "by_type": {
                "flashcard": len(materials.flashcards),
                "quiz": len(materials.quizzes),
                "mnemonic": len(materials.mnemonics)
            }
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/health", methods=["GET"])
def api_health():
    """Comprehensive health check including LLM providers"""
    try:
        # Check LLM health
        llm = get_local_llm()
        llm_healthy, llm_message, llm_details = llm.health_check()
        
        # Check basic app health
        db = load_db()
        card_count = len(db.get("cards", []))
        
        # Check AI pipeline
        pipeline_healthy = True
        pipeline_error = None
        try:
            pipeline = get_ai_pipeline()
        except Exception as e:
            pipeline_healthy = False
            pipeline_error = str(e)
        
        health_status = {
            "status": "healthy" if llm_healthy and pipeline_healthy else "degraded",
            "timestamp": datetime.now().isoformat(),
            "components": {
                "llm": {
                    "status": "healthy" if llm_healthy else "unhealthy",
                    "message": llm_message,
                    "details": llm_details
                },
                "database": {
                    "status": "healthy",
                    "card_count": card_count
                },
                "ai_pipeline": {
                    "status": "healthy" if pipeline_healthy else "unhealthy",
                    "error": pipeline_error
                }
            }
        }
        
        status_code = 200 if (llm_healthy and pipeline_healthy) else 503
        return jsonify(health_status), status_code
        
    except Exception as e:
        return jsonify({
            "status": "unhealthy",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }), 500


# Handle favicon and apple touch icons to avoid 404 noise
@app.route("/favicon.ico")
def favicon():
    try:
        return send_file("static/favicon.ico", mimetype="image/vnd.microsoft.icon")
    except FileNotFoundError:
        # Return empty response if favicon doesn't exist
        from flask import make_response
        response = make_response('', 404)
        response.headers['Content-Type'] = 'image/vnd.microsoft.icon'
        return response


@app.route("/apple-touch-icon.png")
def apple_touch_icon():
    try:
        return send_file("static/apple-touch-icon.png", mimetype="image/png")
    except FileNotFoundError:
        # Return empty response if icon doesn't exist
        from flask import make_response
        response = make_response('', 404)
        response.headers['Content-Type'] = 'image/png'
        return response


@app.route("/")
def index():
    return render_template("index.html")


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=int(os.getenv("PORT", 5000)), debug=True)
